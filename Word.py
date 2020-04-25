import docx
import os
os.chdir('C:\\Users\koenl\Dropbox\Documents\Coding\Python\Automate the Boring Stuff')

d = docx.Document('C:\\Users\koenl\Dropbox\Documents\Coding\Python\Automate the Boring Stuff\demo.docx')
print(d.paragraphs)

d.paragraphs[0]
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]
p.runs
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print(p.runs[1].bold)
print(p.runs[1].italic)

p.runs[3].italic == True
print(p.runs[3].text)
d.save('demo2.docx')

d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('Hello this is another paragraph')
d.save('demo2.docx')

p = d.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True
d.save('demo3.docx')

print('')

def getText(d):
    doc = docx.Document(d)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))
